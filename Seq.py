from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import pickle

FILE_NAME_PREFIX = 'serialized_telo\\'

MIN_TO_ALIGN = 20

EMPTY = ''

INF = float("inf")


class Motif:
    seq = None
    color = None

    def __init__(self):

        self.motif = []
        self.curr_base = 0
        self.motif_error = False
        self.base_error = False
        self.sequence = None

    def add_base(self, base_tuple):
        self.motif.append(base_tuple)
        if base_tuple[0] != self.seq[base_tuple[1]]:
            self.base_error = True
        if self.curr_base != base_tuple[1]:
            self.motif_error = True
        self.curr_base += 1

    def print_doc(self, doc):
        self.generate_seq()
        word = doc.add_run(self.sequence)
        word.bold = self.motif_error or self.base_error
        word.underline = True
        word.font.highlight_color = self.color
        doc.add_run(' ')

    def generate_seq(self):
        if self.sequence:
            return
        self.sequence = ''
        for i in self.motif:
            self.sequence += i[0]


class TTAGGG(Motif):
    seq = 'TTAGGG'
    color = WD_COLOR_INDEX.YELLOW


class TTAAAA(Motif):
    seq = 'TTAAAA'
    color = WD_COLOR_INDEX.RED


class TTGGGG(Motif):
    seq = 'TTGGGG'
    color = WD_COLOR_INDEX.BRIGHT_GREEN


class Telomere:
    motifs = [TTAGGG, TTAAAA, TTGGGG]

    def __init__(self):
        self.seq = []
        self.last_motif_stand = INF
        self.curr_in_background = False
        self.curr_in_pre_telo = None
        self.sequence = None
        self.len = None
        self.motif_types_num = [0, 0, 0]
        self.num_of_motifs = 0
        self.num_of_motifs_errors = None
        self.motif_dict = {}

    def add_motif_base(self, motif_type, base_tuple):
        if base_tuple[1] < self.last_motif_stand or self.curr_in_background:
            self.seq.append(self.motifs[motif_type]())
            self.motif_types_num[motif_type] += 1
            self.num_of_motifs += 1
            self.last_motif_stand = 0
        self.seq[-1].add_base(base_tuple)
        self.last_motif_stand = base_tuple[1]
        self.curr_in_background = False

    def add_background_base(self, base):
        if self.last_motif_stand == 0:
            self.last_motif_stand = INF
        self.curr_in_background = True
        self.seq.append(base)

    def print_statistics(self, doc=None):
        if not self.num_of_motifs_errors:
            self.num_of_motifs_errors = 0
            for elem in self.seq:
                if isinstance(elem, Motif):
                    if elem.base_error or elem.motif_error:
                        self.num_of_motifs_errors += 1
        if self.num_of_motifs > 0:
            if doc:
                doc.add_run(" motif error rate: %.2f" % (self.num_of_motifs_errors / self.num_of_motifs) + "\n")
            else:
                print(" motif error rate: %.2f" % (self.num_of_motifs_errors / self.num_of_motifs))
        else:
            if doc:
                doc.add_run("no motifs, just pre-telo" + "\n")
            else:
                print("no motifs, just pre-telo")
        if doc:
            doc.add_run(" num of motifs:" + str(self.num_of_motifs) + "\n")
        else:
            print(" num of motifs:" + str(self.num_of_motifs))
        for elem in range(len(self.motif_types_num)):
            if self.motif_types_num[elem] > 0:
                percentage = round(self.motif_types_num[elem] / self.num_of_motifs, 2)
            else:
                percentage = 0
            self.motif_types_num[elem] = percentage
        if doc:
            doc.add_run(" motif types division: " + str(self.motif_types_num) + "\n")
        else:
            print(" motif types division: " + str(self.motif_types_num))

    def print_doc(self, doc):
        for k in self.seq:
            if isinstance(k, Motif):
                k.print_doc(doc)
            else:
                doc.add_run(k).underline = True

    def generate_seq(self):
        if self.sequence:
            return
        self.sequence = ''
        for motif in self.seq:
            if isinstance(motif, Motif):
                motif.generate_seq()
                if motif.sequence in self.motif_dict:
                    self.motif_dict[motif.sequence] += 1
                else:
                    self.motif_dict[motif.sequence] = 1
                self.sequence += motif.sequence
            else:
                self.sequence += motif
        self.len = len(self.sequence)


class Seq:

    def __init__(self, length, rec_num):
        self.seq = []
        self.contain_telo = False
        self.more_then_one_telo = False
        self.telo_prefix_error = False
        self.telo_start = None
        self.len = length
        self.rec_num = str(rec_num)
        self.next_telo = None
        self.sequence = None
        self.non_telomeric_parts = None
        self.num_of_motifs_errors = 0
        self.num_of_motifs = 0
        self.longest_telomere_len = 0
        self.motif_dict = {}

    def add_telomere(self):
        if self.next_telo is None:
            self.next_telo = Telomere()
        self.seq.append(self.next_telo)
        self.next_telo = None
        if self.contain_telo:
            self.more_then_one_telo = True
        self.contain_telo = True

    def add_motif_base(self, motif_type, base_tuple, position):
        if not isinstance(self.seq[-1], Telomere):
            self.add_telomere()
        self.seq[-1].add_motif_base(motif_type, base_tuple)
        if self.telo_start is None:
            self.telo_start = position

    def add_telo_background(self, base, position):
        if not isinstance(self.seq[-1], Telomere):
            self.add_telomere()
        self.seq[-1].add_background_base(base)
        if self.telo_start is None:
            self.telo_start = position

    def add_normal_dna_base(self, base_tuple):
        self.seq.append(base_tuple[0])

    def add_pre_telo(self, base_tuple):
        if self.next_telo is None:
            self.next_telo = Telomere()
        self.seq.append(base_tuple[0])

    def print_statistics(self, doc=None, counter=None):
        if self.contain_telo:
            if doc:
                doc.add_run("record %s contain telo, begin at %.2f" % (self.rec_num, self.telo_start / self.len) + "\n")
                doc.add_run("length - " + str(self.len) + "\n")
            else:
                print("record %s contain telo, begin at %.2f" % (self.rec_num, self.telo_start / self.len))
                print("length - " + str(self.len))
        else:
            if counter % 50 == 0:
                print('No telo ' + self.rec_num + '\n')
        if self.more_then_one_telo:
            if doc:
                doc.add_run("2 telo or more" + "\n")
            else:
                print("2 telo or more")
        if self.telo_prefix_error:
            if doc:
                doc.add_run("prefix error" + "\n")
            else:
                print("prefix error")
        telo_index = 0
        for elem in self.seq:
            if isinstance(elem, Telomere):
                if doc:
                    doc.add_run("telo " + str(telo_index) + "\n" + "\n")
                else:
                    print("telo " + str(telo_index) + "\n")
                elem.print_statistics(doc)
                telo_index += 1
                self.num_of_motifs_errors += elem.num_of_motifs_errors
                self.num_of_motifs += elem.num_of_motifs

    def print_doc(self):
        document = Document()
        doc = document.add_paragraph('')
        self.print_statistics(doc)
        for elem in self.seq:
            if isinstance(elem, Telomere):
                elem.print_doc(doc)
            else:
                doc.add_run(elem)
        document.save(self.rec_num + '.docx')

    def generate_seq(self):
        if self.sequence:
            return
        self.sequence = EMPTY
        self.non_telomeric_parts = []
        curr_seq_to_align = EMPTY
        for elem in self.seq:
            if isinstance(elem, Telomere):
                elem.generate_seq()
                for motif in elem.motif_dict:
                    if motif in self.motif_dict:
                        self.motif_dict[motif] += elem.motif_dict[motif]
                    else:
                        self.motif_dict[motif] = elem.motif_dict[motif]
                self.longest_telomere_len = max(self.longest_telomere_len, elem.len)
                self.sequence += elem.sequence
                if curr_seq_to_align != EMPTY:
                    if len(curr_seq_to_align) > MIN_TO_ALIGN:
                        self.non_telomeric_parts.append(curr_seq_to_align)
                    curr_seq_to_align = EMPTY
            else:
                self.sequence += elem
                curr_seq_to_align += elem
        if curr_seq_to_align != EMPTY:
            self.non_telomeric_parts.append(curr_seq_to_align)

    def save_to_file(self):
        if self.contain_telo:
            self.generate_seq()
            file_handler = open(FILE_NAME_PREFIX + self.rec_num + '.obj', 'wb')
            pickle.dump(self, file_handler)
            file_handler.close()
